import sys
import shutil
from tqdm import tqdm
import requests
import time
from requests.utils import dict_from_cookiejar, cookiejar_from_dict
import json
import re
import string
from docx import Document
import os
from docx.oxml.ns import qn
from docx.shared import RGBColor
from ffmpy import FFmpeg


class mooc_login():

    def __init__(self):
        self.session = requests.session()
        self.session.headers = {
            'referer': 'https://www.icourse163.org/',
            'sec-fetch-dest': 'empty',
            'sec-fetch-mode': 'cors',
            'sec-fetch-site': 'same-origin',
            'user-agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/88.0.4324.190 Safari/537.36',

        }

    def get_codeUrl_pollKey(self):
        self.session.get("https://www.icourse163.org/")
        url = "https://www.icourse163.org/logonByQRCode/code.do?width=182&height=182"
        response = self.session.get(url=url)
        dic = response.json()
        if dic["code"] == 0:
            print("获取codeUrl与pollKey成功")
        else:
            print("获取codeUrl与pollKey失败")
        codeUrl = dic["result"]["codeUrl"]
        pollKey = dic["result"]["pollKey"]
        return codeUrl, pollKey

    def show_QR_img(self, codeUrl):
        response = self.session.get(url=codeUrl)
        content = response.content
        with open('core\QR.png', 'wb') as f:
            f.write(content)
        os.startfile('core\QR.png')

        # barcode_url = ''
        # barcodes = decode(Image.open('core/QR.png'))
        # for barcode in barcodes:
        #     barcode_url = barcode.data.decode("utf-8")
        # # print(barcode_url)
        # qr = qrcode.QRCode()
        # qr.add_data(barcode_url)
        # # invert=True白底黑块,有些app不识别黑底白块.
        # qr.print_ascii(invert=True)

    def get_status(self, pollKey):
        url = 'https://www.icourse163.org/logonByQRCode/poll.do?pollKey={}'.format(pollKey)
        while True:
            response = self.session.get(url=url)
            dic = response.json()
            if dic["result"]["codeStatus"] == 0:
                print("请及时扫码")
            elif dic["result"]["codeStatus"] == 1:
                print("请点击确认")
            else:
                print("登录成功")
                token = dic["result"]["token"]
                return token
            time.sleep(1)

    def save_cookie(self):
        codeUrl, pollKey = self.get_codeUrl_pollKey()
        self.show_QR_img(codeUrl)
        token = self.get_status(pollKey)
        params = {
            "token": token,
            "returnUrl": "aHR0cHM6Ly93d3cuaWNvdXJzZTE2My5vcmcv",
        }
        url = "https://www.icourse163.org/passport/logingate/mocMobChangeCookie.htm"
        self.session.get(url=url, params=params)
        self.session.get("https://www.icourse163.org/")
        cookie_jar = self.session.cookies
        cookie_dic = dict_from_cookiejar(cookie_jar)
        with open('core/cookies.json', 'w') as f:
            f.write(json.dumps(cookie_dic))
        print("cookie已保存")

    def check_cookie(self):
        print('正在检查cookie有效性')
        url = 'https://www.icourse163.org/'
        response = self.session.get(url=url)
        text = response.text
        if "个人中心" in text:
            print("cookie有效")
            return True
        else:
            print("cookie无效")
            return False

    def reade_cookie(self):
        b = os.path.exists("core/cookies.json")
        if b == False:
            print("没有cookie文件")
            return False
        else:
            with open(r'core/cookies.json', 'r') as f:
                cookie_dic = json.loads(f.read())
            cookie_jar = cookiejar_from_dict(cookie_dic)
            self.session.cookies = cookie_jar
            print("cookie已加载")
            return True

    def login(self):
        b = os.path.exists('core')
        if b == False:
            os.mkdir('core')
        b = self.reade_cookie()
        if b == True:
            b1 = self.check_cookie()
            if b1 == False:
                # 删除存在的失效cookie
                self.session.cookies.clear()
                self.save_cookie()
                os.remove('core/QR.png')
        else:
            self.save_cookie()
            os.remove('core/QR.png')

        self.session.headers = {
            "sec-ch-ua": '"Google Chrome";v="89", "Chromium";v="89", ";Not A Brand";v="99"',
            "sec-ch-ua-mobile": "?0",
            "sec-fetch-dest": "document",
            "sec-fetch-mode": "navigate",
            "sec-fetch-site": "same-origin",
            "upgrade-insecure-requests": "1",
            "user-agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/89.0.4389.72 Safari/537.3",
        }

        return self.session


class mooc_spider():

    def __init__(self):
        self.session = mooc_login().login()

    def get_csrfKey_userId(self):
        self.userId = dict_from_cookiejar(self.session.cookies)["NETEASE_WDA_UID"].split('#')[0]
        params = {
            "userId": self.userId
        }
        self.session.get('https://www.icourse163.org/home.htm', params=params)
        self.csrfKey = dict_from_cookiejar(self.session.cookies)["NTESSTUDYSI"]
        self.session.headers.update({"edu-script-token": self.csrfKey})
        return

    def get_pagesize(self):
        url = "https://www.icourse163.org/web/j/learnerCourseRpcBean.getMyLearnedCoursePanelList.rpc"
        params = {
            "csrfKey": self.csrfKey
        }
        data = {
            'type': '30',
            'p': '1',
            'psize': '8',
            'courseType': '1',
        }
        response = self.session.post(url=url, params=params, data=data)
        pagesize = response.json()["result"]["pagination"]["totlePageCount"]
        return int(pagesize)

    def get_courses(self, pagesize):
        courses_dic = {}
        url = "https://www.icourse163.org/web/j/learnerCourseRpcBean.getMyLearnedCoursePanelList.rpc"
        params = {
            "csrfKey": self.csrfKey
        }
        n_key = 0
        for p in range(1, pagesize):
            data = {
                "type": "30",
                "p": p,
                "psize": '8',
                "courseType": "1",
            }
            response = self.session.post(url, data=data, params=params)
            result = response.json()["result"]["result"]
            try:
                for i in result:
                    name = i["name"]
                    courses_id = i["termPanel"]["id"]
                    school_name = i["schoolPanel"]["name"]
                    list = [name, school_name, courses_id]
                    courses_dic[str(n_key)] = list
                    n_key += 1
            except:
                pass
        return courses_dic

    def get_timestamp(self):
        timestamp = int(time.time() * 1000)
        return str(timestamp)

    # httpSessionId  NTESSTUDYSI
    def get_NTESSTUDYSI(self):
        NTESSTUDYSI = dict_from_cookiejar(self.session.cookies)["NTESSTUDYSI"]
        return NTESSTUDYSI

    def get_course_text(self, courses_id):
        url = "https://www.icourse163.org/dwr/call/plaincall/CourseBean.getLastLearnedMocTermDto.dwr"
        data = {
            'callCount': '1',
            'scriptSessionId': '${scriptSessionId}190',
            'httpSessionId': self.get_NTESSTUDYSI(),
            'c0-scriptName': 'CourseBean',
            'c0-methodName': 'getLastLearnedMocTermDto',
            'c0-id': '0',
            'c0-param0': 'number:{}'.format(courses_id),
            'batchId': self.get_timestamp()
        }
        response = self.session.post(url=url, data=data)
        text = response.text.encode('utf8').decode('unicode-escape')
        return text

    '''
    下面是测验爬取
    '''

    # ------------------------------------------------------------------------------------------------
    # 获取测验相关信息
    def get_test_info(self, text):
        test_name_id_dic = {}
        # id_name_list = re.findall(
        #     's\d{1,3}\.id=(\d+);s\d{1,3}\.name="(.*?测试.*?)";|s\d{1,3}\.id=(\d+);s\d{1,3}\.name="(.*?测验.*?)";', text)
        id_name_list = re.findall(
            's\d{1,3}\.id=(\d+);s\d{1,3}\.name="(.*?)";', text)
        for id_name in id_name_list:
            test_id = id_name[0]
            test_name = id_name[1]
            # for id_name in id_name_list:
            #     if id_name[0] != '' and id_name[1] != '':
            #         test_id = id_name[0]
            #         test_name = id_name[1]
            #     else:
            #         test_id = id_name[2]
            #         test_name = id_name[3]
            data = {
                'callCount': '1',
                'scriptSessionId': '${scriptSessionId}190',
                'httpSessionId': self.get_NTESSTUDYSI(),
                'c0-scriptName': 'MocQuizBean',
                'c0-methodName': 'getQuizPaperDto',
                'c0-id': '0',
                'c0-param0': f'string:{test_id}',
                'c0-param1': 'number:0',
                'c0-param2': 'boolean:false',
                'batchId': '1619443609960',
            }
            url = 'https://www.icourse163.org/dwr/call/plaincall/MocQuizBean.getQuizPaperDto.dwr'
            response = self.session.post(url=url, data=data)
            text = response.text
            aid = re.findall('{aid:(\d+),', text)
            test_name_id_dic[test_name] = aid
            time.sleep(1.5)
        return test_name_id_dic

    def save_img(self, url):
        headers = {
            'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/88.0.4324.150 Safari/537.36',
        }
        response = requests.get(url=url, headers=headers)
        content = response.content
        with open('core\paper.png', 'wb') as f:
            f.write(content)
        return

    def get_paper(self, aid):
        url = "https://www.icourse163.org/mob/course/paperDetail/v1"
        data = {
            "mob-token": "d6dbe470975c1e411442686996218e49e2a8d084dd67e087e19a94a9b58fd537e0aaa719763eaa7cf3f592bf43a75665af8a8a488536a10030490cd75e4c93b1935c300b4df0b319e73756c2bf4f77453fe82d8881bd18d30e2ebd9980e23f4b0d7ef1ee30366d043e17db84286433cea05f7dfa572380889dc5f9487881ea827610539dc81b92fdbdf29d78a6ea65037f9085857f5ac3a9dcf8392a8ff2c17c8676b27e0d374f2f0db184284eeb568ef068144cfcbaf5b449620a83b20dcbf5",
            "testId": "2222", "isExercise": "false", "aId": aid, "withStdAnswerAndAnalyse": "true"}
        headers = {"Content-Type": "application/x-www-form-urlencoded; charset=UTF-8",
                   'user-agent': "Dalvik/2.1.0 (Linux; U; Android 8.1.0; Pixel Build/OPM4.171019.021.P1)", }
        response = requests.post(url, data=data, headers=headers)
        objective_dic = {}
        try:
            objectiveQList = response.json()['results']['mocPaperDto']['objectiveQList']
            for i in range(len(objectiveQList)):
                objective = objectiveQList[i]
                plainTextTitle = objective['plainTextTitle']
                if '【图片】' in plainTextTitle:
                    title = objective['title']
                    src_list = re.findall('src="(.*?)"', title)
                    for m in range(len(src_list)):
                        plainTextTitle = re.sub('【图片】', ',{},'.format(src_list[m]), plainTextTitle, count=1)
                optionDtos_list = objective['optionDtos']
                str = string.ascii_uppercase
                optionDtos_dic = {}
                answer = ''
                if optionDtos_list != []:
                    for n in range(len(optionDtos_list)):
                        optionDtos = optionDtos_list[n]
                        content = optionDtos['content']
                        content = re.sub('<span style.*?>|"|<p>|</p>|</span>|&nbsp;|<br >|', '', content)
                        content = re.sub('<span style=font-size:\d+px;font-family:\n;  >', '', content, re.S)
                        if '<img' in content:
                            src_list = re.findall('src=(.*?)/>', content)
                            for m in range(len(src_list)):
                                content = re.sub('<img .*?src=.*?/>|" "', ',{},'.format(src_list[m]), content, count=1)
                        optionDtos_dic[str[n]] = content.split(',')
                        if optionDtos['answer'] == True:
                            answer = str[n]
                        n += 1
                else:
                    answer = objective['stdAnswer']
                num = i + 1
                objective_dic[num] = {}
                objective_dic[num]['plainTextTitle'] = plainTextTitle.split(',')
                objective_dic[num]['option'] = optionDtos_dic
                objective_dic[num]['answer'] = answer
        except:
            pass
        return objective_dic

    def word(self, objective_dic, path, test_name):
        # 试卷document
        document = Document()
        document.styles['Normal'].font.name = u'宋体'
        document.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
        document.styles['Normal'].font.color.rgb = RGBColor(0, 0, 0)
        for title_num, objective in tqdm(objective_dic.items(), desc=test_name):
            #     '''开启一个新段落'''
            paragraph = document.add_paragraph('{},'.format(title_num))
            # 设置行距
            paragraph_format = paragraph.paragraph_format
            paragraph_format.line_spacing = 1.5  # 1.5倍行距
            run = document.paragraphs[-1].add_run()
            # 题目信息
            plainTextTitle = objective['plainTextTitle']
            for i in plainTextTitle:
                try:
                    if 'http' in i:
                        url = i.replace('  ', '')
                        self.save_img(url)
                        run.add_picture('core\paper.png')
                    elif '//nos.netease.com/edu-image' in i:
                        url = ('http:' + i).replace('  ', '')
                        self.save_img(url)
                        run.add_picture('core\paper.png')
                    else:
                        run.add_text(u'{}'.format(i))
                except:
                    pass
            # 选项
            option = objective['option']
            if option != {}:
                for key, value in option.items():
                    '''key代表选项ABCD,value代表选项内容，是一个列表'''
                    '''开启一个新段落'''
                    paragraph = document.add_paragraph('{},'.format(key))
                    # 设置行距
                    paragraph_format = paragraph.paragraph_format
                    paragraph_format.line_spacing = 1.5  # 1.5倍行距
                    run = document.paragraphs[-1].add_run()
                    for m in value:
                        try:
                            if 'http' in m:
                                url = m.replace('  ', '')
                                self.save_img(url)
                                run.add_picture('core\paper.png')
                            elif '//nos.netease.com/edu-image' in m:
                                url = ('http:' + m).replace('  ', '')
                                self.save_img(url)
                                run.add_picture('core\paper.png')
                            else:
                                run.add_text(u'{}'.format(m))
                        except:
                            pass
            else:
                pass
            '''答案'''
            answer = objective['answer']
            '''开启一个新段落'''
            paragraph = document.add_paragraph('答案:{}'.format(answer))
            # 设置行距
            paragraph_format = paragraph.paragraph_format
            paragraph_format.line_spacing = 1.5  # 1.5倍行距
            document.save('{}/{}.docx'.format(path, test_name))

    # -------------------------------------------------------------------------------------------------
    def get_info_name_id(self, text):
        """
        获取视频等资料的，名字与id
        :param text: text为从返回的信息中提取我们需要的参数
        :return:dict_document, other_dict, dict_movie,xmind分别为ppt的名字与id，ppt的名字与id，其它资料的名字与id，视频的名字与id，
                        xmind为整个课程的逻辑，讨论，作业，不在里面显示
        """
        lis = re.findall('s1\[\d+]=(s\d+)', text)
        num = len(lis)
        dict_document = {}
        dict_movie = {}
        other_dict = {}
        xmind = ""
        for i in range(num):
            if i == num - 1:
                chapter = re.findall(f'{lis[i]}\.contentId.*?dwr\.engine', text, re.S)[0]
                chapter_name = re.findall('lessons=s\d+;s\d+\.name="(.*?)";', chapter)[0]
                xmind += f"|--->>{chapter_name}\n"
                lis1 = re.findall(
                    'contentId=(\d+);s\d+\.contentType=([0-4]|[7-9]);.*?s\d+.id=(\d+).*?name="(.*?)";.*?yktRelatedLiveInfo',
                    chapter)
                for i in lis1:
                    contentId, contentType, id, name = i
                    if int(contentType) == 1:
                        dict_movie[name] = {}
                        dict_movie[name]['contentId'] = contentId
                        dict_movie[name]['id'] = id
                    elif int(contentType) == 3:
                        dict_document[name] = {}
                        dict_document[name]['contentId'] = contentId
                        dict_document[name]['id'] = id
                    else:
                        other_dict[name] = {}
                        other_dict[name]['contentId'] = contentId
                        other_dict[name]['id'] = id
                    xmind += f"|----------{name}\n"
            else:

                chapter = re.findall(f'{lis[i]}\.contentId.*?{lis[i + 1]}\.contentId', text, re.S)[0]
                chapter_name = re.findall('lessons=s\d+;s\d+\.name="(.*?)";', chapter)[0]
                xmind += f"|--->>{chapter_name}\n"
                lis1 = re.findall(
                    'contentId=(\d+);s\d+\.contentType=([0-4]|[7-9]);.*?s\d+.id=(\d+).*?name="(.*?)";.*?yktRelatedLiveInfo',
                    chapter)
                for i in lis1:
                    contentId, contentType, id, name = i
                    if int(contentType) == 1:
                        dict_movie[name] = {}
                        dict_movie[name]['contentId'] = contentId
                        dict_movie[name]['id'] = id
                    elif int(contentType) == 3:
                        dict_document[name] = {}
                        dict_document[name]['contentId'] = contentId
                        dict_document[name]['id'] = id
                    else:
                        other_dict[name] = {}
                        other_dict[name]['contentId'] = contentId
                        other_dict[name]['id'] = id
                    xmind += f"|----------{name}\n"
        print(xmind)

        return dict_document, other_dict, dict_movie

    def download_ppt(self, dict_document, path):
        url = 'https://www.icourse163.org/dwr/call/plaincall/CourseBean.getLessonUnitLearnVo.dwr'
        for name, value in dict_document.items():
            contentId = value['contentId']
            id = value['id']
            payload = {
                'callCount': '1',
                'scriptSessionId': '${scriptSessionId}190',
                'httpSessionId': self.get_NTESSTUDYSI(),
                'c0-scriptName': 'CourseBean',
                'c0-methodName': 'getLessonUnitLearnVo',
                'c0-id': '0',
                'c0-param0': f'number:{contentId}',
                'c0-param1': 'number:3',
                'c0-param2': 'number:0',
                'c0-param3': f'number:{id}',
                'batchId': self.get_timestamp(),
            }
            response = self.session.post(url=url, data=payload)
            text = response.text
            ppt_url = re.findall('textUrl:"(.*?)",', text)[0]
            response = requests.get(url=ppt_url)
            temp_size = 0  # 已经下载文件大小
            chunk_size = 1024  # 每次下载数据大小
            start = time.time()
            total_size = int(response.headers.get("Content-Length"))
            with open(fr'{path}\{name}.pdf', 'wb') as f:
                for chunk in response.iter_content(chunk_size=chunk_size):
                    if chunk:
                        temp_size += len(chunk)
                        f.write(chunk)
                        f.flush()
                        #############花哨的下载进度部分###############
                        done = int(50 * temp_size / total_size)
                        # 调用标准输出刷新命令行，看到\r 回车符了吧
                        # 相当于把每一行重新刷新一遍
                        sys.stdout.write(
                            "\r[%s%s] %d%%" % ('█' * done, ' ' * (50 - done), 100 * temp_size / total_size))
                        sys.stdout.flush()
            print()  # 避免上面\r 回车符，执行完后需要换行了，不然都在一行显示
            end = time.time()  # 结束时间
            print(f'{name}下载完成!用时%.2f 秒' % (end - start))

    def handle_ts(self, text, y, m, d, name, save_path):
        if os.path.exists('ts') == False:
            os.mkdir('ts')
        ts_list = re.findall('.*?ts', text)
        headers = {
            'user-agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/90.0.4430.85 Safari/537.36',
        }
        num = len(ts_list)
        path_list = [fr'ts\00{i}.ts' for i in range(num)]
        for i in range(num):
            ts = ts_list[i]
            url = f'https://mooc2vod.stu.126.net/nos/hls/{y}/{m}/{d}/{ts}'
            response = requests.get(url=url, headers=headers)
            path = fr'ts\00{i}.ts'
            temp_size = 0  # 已经下载文件大小
            chunk_size = 1024  # 每次下载数据大小
            start = time.time()
            total_size = int(response.headers.get("Content-Length"))
            with open(path, 'wb') as f:
                for chunk in response.iter_content(chunk_size=chunk_size):
                    if chunk:
                        temp_size += len(chunk)
                        f.write(chunk)
                        f.flush()
                        #############花哨的下载进度部分###############
                        done = int(50 * temp_size / total_size)
                        # 调用标准输出刷新命令行，看到\r 回车符了吧
                        # 相当于把每一行重新刷新一遍
                        sys.stdout.write(
                            "\r[%s%s] %d%%" % ('█' * done, ' ' * (50 - done), 100 * temp_size / total_size))
                        sys.stdout.flush()
            print()  # 避免上面\r 回车符，执行完后需要换行了，不然都在一行显示
            end = time.time()  # 结束时间
            print(f'ts{i+1}下载完成!用时%.2f 秒' % (end - start))

        ff = FFmpeg(inputs={'concat:' + '|'.join(path_list): None},
                    outputs={fr'{save_path}\{name}': '-loglevel quiet -c copy -absf aac_adtstoasc -movflags faststart'})
        print(f'----------》》》{name}完成!')
        ff.run()
        shutil.rmtree('ts')
        return

    def download_movie(self, dict_movie, path, Video_quality):
        url = 'https://www.icourse163.org/web/j/resourceRpcBean.getResourceToken.rpc'
        params = {
            'csrfKey': self.get_NTESSTUDYSI()
        }
        for name, value in dict_movie.items():
            id = value['id']
            data = {
                'bizId': id,
                'bizType': '1',
                'contentType': '1',
            }
            response = self.session.post(url=url, params=params, data=data)
            result = response.json()['result']
            videoSignDto = result['videoSignDto']
            signature = videoSignDto['signature']
            videoId = videoSignDto['videoId']

            def movie(signature, videoId):
                url = 'https://vod.study.163.com/eds/api/v1/vod/video'
                params = {
                    'videoId': videoId,
                    'signature': signature,
                    'clientType': '1',
                }
                response = self.session.post(url=url, params=params)
                result = response.json()['result']
                name = result['name']
                videos = result['videos']
                videoUrl = videos[Video_quality]['videoUrl']
                y, m, d = re.findall('http://mooc2vod.stu.126.net/nos/hls/(\d+)/(\d+)/(\d+)', videoUrl)[0]
                headers = {
                    'user-agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/90.0.4430.85 Safari/537.36',
                }
                response = requests.get(url=videoUrl, headers=headers)
                text = response.text
                if '.mp4' not in name:
                    name += '.mp4'
                if os.path.exists('ts') == True:
                    shutil.rmtree('ts')
                if os.path.exists(fr'{path}\{name}') == True:
                    print(fr'{name}已经存在!')
                    return
                else:
                    self.handle_ts(text, y, m, d, name, path)

            movie(signature, videoId)

    # -------------------------------------------------------------------------------------------------

    def spider(self):
        b = os.path.exists('data')
        if b == False:
            os.mkdir('data')
        self.get_csrfKey_userId()
        pagessize = self.get_pagesize()
        courses_dic = self.get_courses(pagessize)
        for key, value in courses_dic.items():
            print(key, value[0], value[1])
        while True:
            print('-' * 100)
            key = input("输入非纯字符将会自动退出\n请输入你要选择课程的序号:")
            if key.isdigit() == True:
                path = 'data/{}_{}'.format(courses_dic[key][0], courses_dic[key][1])
                b = os.path.exists(path)
                if b == False:
                    os.mkdir(path)
                courses_id = courses_dic[key][2]
                text = self.get_course_text(courses_id)
                dict_document, other_dict, dict_movie = self.get_info_name_id(text)
                print('----------------------------')
                print('0----》试卷\n'
                      '1----》视频\n'
                      '2----》pdf等资料\n'
                      '3----》以上全部\n'
                      '其它代表重新选择课程')

                def paper():
                    print('-' * 100)
                    print('正在爬取 {} 课程 {} 测验'.format(courses_dic[key][1], courses_dic[key][0]))
                    test_name_id_dic = self.get_test_info(text)
                    print('开始下载......')
                    for test_name, aid in test_name_id_dic.items():
                        objective_dic = self.get_paper(aid)
                        if objective_dic != {}:
                            self.word(objective_dic, path, test_name)

                while True:
                    num = input('请输入:')
                    if num == '0':
                        paper()
                    elif num == '1':
                        Video_quality = input('0---》标清\n'
                                              '1---》高清\n'
                                              '2---》超清\n'
                                              '请输入视频清晰度:')
                        Video_quality = int(Video_quality)
                        self.download_movie(dict_movie, path, Video_quality)
                    elif num == '2':
                        self.download_ppt(dict_document, path)
                    elif num == '3':
                        paper()
                        Video_quality = input('0---》标清\n'
                                              '1---》高清\n'
                                              '2---》超清\n'
                                              '请输入视频清晰度:')
                        Video_quality = int(Video_quality)
                        self.download_movie(dict_movie, path, Video_quality)
                        self.download_ppt(dict_document, path)
                    else:
                        break

            else:
                sys.exit()


m = mooc_spider()
m.spider()
