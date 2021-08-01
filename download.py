# -*- encoding: utf-8 -*-
"""
@File    :   spider.py    
@Contact :   jyj345559953@qq.com
@Author  :   Esword
"""
import os
import re
import shutil
import string
import sys
import time
import requests
from docx import Document
from docx.oxml.ns import qn
from docx.shared import RGBColor
from ffmpy import FFmpeg
from requests.utils import dict_from_cookiejar
from tqdm import tqdm


class Download():

    def __init__(self, session):
        self.session = session
        self.session.headers = {
            "user-agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/92.0.4515.107 Safari/537.36 Edg/92.0.902.55",
        }
        self.headers = {
            "user-agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/91.0.4472.164 Safari/537.36"
        }

    def get_csrfKey_userId(self):
        self.userId = dict_from_cookiejar(self.session.cookies)["NETEASE_WDA_UID"].split('#')[0]
        params = {
            "userId": self.userId
        }
        self.session.get('https://www.icourse163.org/home.htm', params=params)
        self.csrfKey = dict_from_cookiejar(self.session.cookies)["NTESSTUDYSI"]
        self.session.headers.update({"edu-script-token": self.csrfKey})
        return

    def get_timestamp(self):
        timestamp = int(time.time() * 1000)
        return str(timestamp)

    # httpSessionId  NTESSTUDYSI
    def get_NTESSTUDYSI(self):
        NTESSTUDYSI = dict_from_cookiejar(self.session.cookies)["NTESSTUDYSI"]
        return NTESSTUDYSI

    # 这里是对课程的信息处理及分类
    def get_info_and_handle_info(self, course_id):
        url = f"https://www.icourse163.org/web/j/courseBean.getLastLearnedMocTermDto.rpc?csrfKey={self.csrfKey}"
        data = {
            "termId": course_id
        }
        response = self.session.post(url=url, data=data)
        mocTermDto = response.json()["result"]["mocTermDto"]
        Data = {}
        chapters = mocTermDto["chapters"]
        for chapter in chapters:
            chaptername = chapter["name"]
            contentType = chapter["contentType"]
            Data[chaptername] = {}
            if contentType == 2:
                Data[chaptername]["contentType"] = contentType
                Data[chaptername]["exam"] = {}
                name = chapter["name"]
                id = chapter["id"]
                Data[chaptername]["exam"]["name"] = name
                Data[chaptername]["exam"]["id"] = id
            else:
                Data[chaptername]["contentType"] = contentType
                lessons = chapter["lessons"]
                quizs = chapter["quizs"]
                if lessons != None or lessons != []:
                    for lesson in lessons:
                        lessonname = lesson["name"]
                        Data[chaptername][lessonname] = {}
                        units = lesson["units"]
                        for unit in units:
                            unitname = unit["name"]
                            Data[chaptername][lessonname][unitname] = {}
                            id = unit["id"]
                            name = unit["name"]
                            contentType = unit["contentType"]
                            contentId = unit["contentId"]
                            Data[chaptername][lessonname][unitname]["name"] = name
                            Data[chaptername][lessonname][unitname]["id"] = id
                            Data[chaptername][lessonname][unitname]["contentType"] = contentType
                            if contentType == 3 or contentType == 4 or contentType == 5 or contentType == 6:
                                Data[chaptername][lessonname][unitname]["contentId"] = contentId

                if quizs == None or quizs == []:
                    pass
                else:
                    quizs = quizs[0]
                    contentType = quizs["contentType"]
                    quizname = quizs["name"]
                    contentId = quizs["contentId"]
                    Data[chaptername][quizname] = {}
                    Data[chaptername][quizname]["contentType"] = contentType
                    Data[chaptername][quizname]["contentId"] = contentId

        return Data

    # 获取movie下载的 name, signature, videoId
    def Getmovieurl(self, id):
        data = {
            "bizId": id,
            "bizType": "1",
            "contentType": "1",
        }
        url = f"https://www.icourse163.org/web/j/resourceRpcBean.getResourceToken.rpc?csrfKey={self.csrfKey}"
        response = self.session.post(url=url, data=data, headers=self.headers)
        videoSignDto = response.json()["result"]["videoSignDto"]
        name = videoSignDto["name"]
        signature = videoSignDto["signature"]
        videoId = videoSignDto["videoId"]

        params = {
            "videoId": videoId,
            "signature": signature,
            "clientType": 1,
        }
        url = "https://vod.study.163.com/eds/api/v1/vod/video"
        response = requests.get(url=url, params=params, headers=self.headers)
        videos = response.json()["result"]["videos"]
        num = len(videos)
        if num == 1:
            videoUrl = videos[-1]["videoUrl"]
            baseurl = re.findall('(http.*?/nos/hls/.*/)', videoUrl)[0]
            return baseurl, videoUrl

        elif (num == 2 and self.quality == 2) or (num == 2 and self.quality == 3):
            videoUrl = videos[-1]["videoUrl"]
            baseurl = re.findall('(http.*?/nos/hls/.*/)', videoUrl)[0]
            return baseurl, videoUrl

        else:
            for video in videos:
                quality = video["quality"]
                if quality == self.quality:
                    videoUrl = video["videoUrl"]
                    baseurl = re.findall('(http.*?/nos/hls/.*/)', videoUrl)[0]
                    return baseurl, videoUrl

    # 下载ts文件，并合成视频
    def download_(self, tsurl_list, Path, name):
        path_list = []
        i = 0
        for tsurl in tqdm(tsurl_list):
            path = fr'tem\00{i}.ts'
            i += 1
            path_list.append(path)
            response = requests.get(url=tsurl, headers=self.headers)
            with open(path, 'wb') as f:
                f.write(response.content)
        ff = FFmpeg(inputs={'concat:' + '|'.join(path_list): None},
                    outputs={fr'{Path}\{name}.mp4': '-loglevel quiet -c copy -absf aac_adtstoasc -movflags faststart'})
        print(f'----------》》》{name}合并完成!    [mp4]\n'
              f'------------------------------------------------')
        ff.run()
        shutil.rmtree('tem')

    def movie(self, id, path, name):
        if os.path.exists("tem") == True:
            shutil.rmtree('tem')
            os.mkdir("tem")
        else:
            os.mkdir("tem")
        baseurl, videoUrl = self.Getmovieurl(id)
        response = requests.get(url=videoUrl, headers=self.headers)
        text = response.text
        ts_list = re.findall('.*?ts', text)
        tsurl_list = [baseurl + ts for ts in ts_list]
        self.download_(tsurl_list, path, name)

    def save_img(self, url):
        headers = {
            'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/88.0.4324.150 Safari/537.36',
        }
        response = requests.get(url=url, headers=headers)
        content = response.content
        with open('core\paper.png', 'wb') as f:
            f.write(content)
        return

    def paper(self, aid):
        url = "https://www.icourse163.org/mob/course/paperDetail/v1"
        data = {
            "mob-token": "d6dbe470975c1e411442686996218e49e2a8d084dd67e087e19a94a9b58fd537e0aaa719763eaa7cf3f592bf43a75665af8a8a488536a10030490cd75e4c93b1935c300b4df0b319e73756c2bf4f77453fe82d8881bd18d30e2ebd9980e23f4b0d7ef1ee30366d043e17db84286433cea05f7dfa572380889dc5f9487881ea827610539dc81b92fdbdf29d78a6ea65037f9085857f5ac3a9dcf8392a8ff2c17c8676b27e0d374f2f0db184284eeb568ef068144cfcbaf5b449620a83b20dcbf5",
            "testId": "2222", "isExercise": "false", "aId": aid, "withStdAnswerAndAnalyse": "true"}
        headers = {"Content-Type": "application/x-www-form-urlencoded; charset=UTF-8",
                   'user-agent': "Dalvik/2.1.0 (Linux; U; Android 8.1.0; Pixel Build/OPM4.171019.021.P1)", }
        response = requests.post(url, data=data, headers=headers)
        objective_dic = {}
        try:
            objectiveQList = response.json()['results']['mocPaperDto']['objectiveQList']
            for i in range(len(objectiveQList)):
                objective = objectiveQList[i]
                plainTextTitle = objective['plainTextTitle']
                if '【图片】' in plainTextTitle:
                    title = objective['title']
                    src_list = re.findall('src="(.*?)"', title)
                    for m in range(len(src_list)):
                        plainTextTitle = re.sub('【图片】', ',{},'.format(src_list[m]), plainTextTitle, count=1)
                optionDtos_list = objective['optionDtos']
                str = string.ascii_uppercase
                optionDtos_dic = {}
                answer = ''
                if optionDtos_list != []:
                    for n in range(len(optionDtos_list)):
                        optionDtos = optionDtos_list[n]
                        content = optionDtos['content']
                        content = re.sub('<span style.*?>|"|<p>|</p>|</span>|&nbsp;|<br >|', '', content)
                        content = re.sub('<span style=font-size:\d+px;font-family:\n;  >', '', content, re.S)
                        if '<img' in content:
                            src_list = re.findall('src=(.*?)/>', content)
                            for m in range(len(src_list)):
                                content = re.sub('<img .*?src=.*?/>|" "', ',{},'.format(src_list[m]), content, count=1)
                        optionDtos_dic[str[n]] = content.split(',')
                        if optionDtos['answer'] == True:
                            answer = str[n]
                        n += 1
                else:
                    answer = objective['stdAnswer']
                num = i + 1
                objective_dic[num] = {}
                objective_dic[num]['plainTextTitle'] = plainTextTitle.split(',')
                objective_dic[num]['option'] = optionDtos_dic
                objective_dic[num]['answer'] = answer
        except:
            pass
        return objective_dic

    def word(self, objective_dic, path, name, quizsbool=False, index=None):
        # 试卷document
        document = Document()
        document.styles['Normal'].font.name = u'宋体'
        document.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
        document.styles['Normal'].font.color.rgb = RGBColor(0, 0, 0)
        for title_num, objective in tqdm(objective_dic.items(), desc=name):
            #     '''开启一个新段落'''
            paragraph = document.add_paragraph('{},'.format(title_num))
            # 设置行距
            paragraph_format = paragraph.paragraph_format
            paragraph_format.line_spacing = 1.5  # 1.5倍行距
            run = document.paragraphs[-1].add_run()
            # 题目信息
            plainTextTitle = objective['plainTextTitle']
            for i in plainTextTitle:
                try:
                    if 'http' in i:
                        url = i.replace('  ', '')
                        self.save_img(url)
                        run.add_picture('core\paper.png')
                    elif '//nos.netease.com/edu-image' in i:
                        url = ('http:' + i).replace('  ', '')
                        self.save_img(url)
                        run.add_picture('core\paper.png')
                    else:
                        run.add_text(u'{}'.format(i))
                except:
                    pass
            # 选项
            option = objective['option']
            if option != {}:
                for key, value in option.items():
                    '''key代表选项ABCD,value代表选项内容，是一个列表'''
                    '''开启一个新段落'''
                    paragraph = document.add_paragraph('{},'.format(key))
                    # 设置行距
                    paragraph_format = paragraph.paragraph_format
                    paragraph_format.line_spacing = 1.5  # 1.5倍行距
                    run = document.paragraphs[-1].add_run()
                    for m in value:
                        try:
                            if 'http' in m:
                                url = m.replace('  ', '')
                                self.save_img(url)
                                run.add_picture('core\paper.png')
                            elif '//nos.netease.com/edu-image' in m:
                                url = ('http:' + m).replace('  ', '')
                                self.save_img(url)
                                run.add_picture('core\paper.png')
                            else:
                                run.add_text(u'{}'.format(m))
                        except:
                            pass
            else:
                pass
            '''答案'''
            answer = objective['answer']
            '''开启一个新段落'''
            paragraph = document.add_paragraph('答案:{}'.format(answer))
            # 设置行距
            paragraph_format = paragraph.paragraph_format
            paragraph_format.line_spacing = 1.5  # 1.5倍行距
            if quizsbool == True:
                document.save(f'{path}\{name}_{index}.docx')
            else:
                document.save(f'{path}\{name}.docx')

    def text(self, path, name, id, contentId):
        url = 'https://www.icourse163.org/dwr/call/plaincall/CourseBean.getLessonUnitLearnVo.dwr'
        payload = {
            'callCount': '1',
            'scriptSessionId': '${scriptSessionId}190',
            'httpSessionId': self.get_NTESSTUDYSI(),
            'c0-scriptName': 'CourseBean',
            'c0-methodName': 'getLessonUnitLearnVo',
            'c0-id': '0',
            'c0-param0': f'number:{contentId}',
            'c0-param1': 'number:4',
            'c0-param2': 'number:0',
            'c0-param3': f'number:{id}',
            'batchId': self.get_timestamp(),
        }
        response = self.session.post(url=url, data=payload)
        text = response.text
        text_url = re.findall('(http.*quality=\d+)', text)[0]
        response = requests.get(url=text_url)
        content = response.content
        with open(fr'{path}\{name}.png', 'wb') as f:
            f.write(content)

    def data(self, path, name, id, contentId):
        url = 'https://www.icourse163.org/dwr/call/plaincall/CourseBean.getLessonUnitLearnVo.dwr'
        payload = {
            'callCount': '1',
            'scriptSessionId': '${scriptSessionId}190',
            'httpSessionId': self.get_NTESSTUDYSI(),
            'c0-scriptName': 'CourseBean',
            'c0-methodName': 'getLessonUnitLearnVo',
            'c0-id': '0',
            'c0-param0': f'number:{contentId}',
            'c0-param1': 'number:3',
            'c0-param2': 'number:0',
            'c0-param3': f'number:{id}',
            'batchId': self.get_timestamp(),
        }
        response = self.session.post(url=url, data=payload)
        text = response.text
        ppt_url = re.findall('textUrl:"(.*?)",', text)[0]
        response = requests.get(url=ppt_url)
        temp_size = 0  # 已经下载文件大小
        chunk_size = 1024  # 每次下载数据大小
        start = time.time()
        total_size = int(response.headers.get("Content-Length"))
        with open(fr'{path}\{name}.pdf', 'wb') as f:
            for chunk in response.iter_content(chunk_size=chunk_size):
                if chunk:
                    temp_size += len(chunk)
                    f.write(chunk)
                    f.flush()
                    #############花哨的下载进度部分###############
                    done = int(50 * temp_size / total_size)
                    # 调用标准输出刷新命令行，看到\r 回车符了吧
                    # 相当于把每一行重新刷新一遍
                    sys.stdout.write(
                        "\r[%s%s] %d%%" % ('█' * done, ' ' * (50 - done), 100 * temp_size / total_size))
                    sys.stdout.flush()
        print()  # 避免上面\r 回车符，执行完后需要换行了，不然都在一行显示
        end = time.time()  # 结束时间
        print(f'{name}下载完成!用时{end - start} 秒    [pdf]\n'
              f'------------------------------------------------')

    def discuss(self, path, name, id, contentId):

        # 获取讨论问题
        def getIssue():
            payload = {
                'callCount': '1',
                'scriptSessionId': '${scriptSessionId}190',
                'httpSessionId': self.get_NTESSTUDYSI(),
                'c0-scriptName': 'CourseBean',
                'c0-methodName': 'getLessonUnitLearnVo',
                'c0-id': '0',
                'c0-param0': f'number:{contentId}',
                'c0-param1': 'number:6',
                'c0-param2': 'number:0',
                'c0-param3': f'number:{id}',
                'batchId': self.get_timestamp(),
            }
            url = "https://www.icourse163.org/dwr/call/plaincall/CourseBean.getLessonUnitLearnVo.dwr"
            response = self.session.post(url=url, data=payload)
            text = response.text.encode('utf-8').decode("unicode_escape")
            issueList = re.findall('content="(.*?)";s\d+\.', text)
            issueList = [re.sub('<.*?>', "", i) for i in issueList]
            titleList = re.findall('s\d+\.title="(.*?)";s\d+\.', text)
            return titleList, issueList

        # 获取回答页数
        def getAnswerpage():
            payload = {
                'callCount': '1',
                'scriptSessionId': '${scriptSessionId}190',
                'httpSessionId': self.get_NTESSTUDYSI(),
                'c0-scriptName': 'PostBean',
                'c0-methodName': 'getPaginationReplys',
                'c0-id': '0',
                'c0-param0': f'number:{contentId}',
                'c0-param1': 'number:2',
                'c0-param2': "number:1",
                'batchId': self.get_timestamp(),
            }
            url = "https://www.icourse163.org/dwr/call/plaincall/PostBean.getPaginationReplys.dwr"
            response = self.session.post(url=url, data=payload)
            text = response.text
            totalPageCount = re.findall('totalPageCount:(\d+),', text)[0]
            return int(totalPageCount)

        # 获取回答
        def getAnswer(totalPageCount):
            AnswerList = []
            url = "https://www.icourse163.org/dwr/call/plaincall/PostBean.getPaginationReplys.dwr"
            for p in range(1, totalPageCount + 1):
                payload = {
                    'callCount': '1',
                    'scriptSessionId': '${scriptSessionId}190',
                    'httpSessionId': self.get_NTESSTUDYSI(),
                    'c0-scriptName': 'PostBean',
                    'c0-methodName': 'getPaginationReplys',
                    'c0-id': '0',
                    'c0-param0': f'number:{contentId}',
                    'c0-param1': 'number:2',
                    'c0-param2': f"number:{p}",
                    'batchId': self.get_timestamp(),
                }
                response = self.session.post(url=url, data=payload)
                text = response.text.encode('utf-8').decode("unicode_escape")
                answerList = re.findall('content="(.*?)";s\d+\.', text)
                answerList = [re.sub('<.*?>', "", i) for i in answerList]
                AnswerList.extend(answerList)
            return AnswerList

        totalPageCount = getAnswerpage()
        titleList, issueList = getIssue()
        AnswerList = getAnswer(totalPageCount)
        with open(f"{path}\{name}.txt", "a+",encoding='utf-8') as f:
            f.write("title:")

        for i in titleList:
            with open(f"{path}\{name}.txt", "a+",encoding='utf-8') as f:
                f.write(i + "\n")

        with open(f"{path}\{name}.txt", "a+",encoding='utf-8') as f:
            f.write("issue:" + "\n")

        for i in issueList:
            with open(f"{path}\{name}.txt", "a+",encoding='utf-8') as f:
                f.write(i + "\n")

        with open(f"{path}\{name}.txt", "a+",encoding='utf-8') as f:
            f.write("answer:" + "\n")

        for i in AnswerList:
            with open(f"{path}\{name}.txt", "a+",encoding='utf-8') as f:
                f.write(i + "\n")

    def test(self, path, name, id, contentId):
        payload = {
            'callCount': '1',
            'scriptSessionId': '${scriptSessionId}190',
            'httpSessionId': self.get_NTESSTUDYSI(),
            'c0-scriptName': 'CourseBean',
            'c0-methodName': 'getLessonUnitLearnVo',
            'c0-id': '0',
            'c0-param0': f'number:{contentId}',
            'c0-param1': 'number:5',
            'c0-param2': 'number:0',
            'c0-param3': f'number:{id}',
            'batchId': self.get_timestamp(),
        }
        url = "https://www.icourse163.org/dwr/call/plaincall/CourseBean.getLessonUnitLearnVo.dwr"
        response = self.session.post(url=url, data=payload)
        text = response.text
        aid = re.findall("aid=(\d+);", text)[0]
        objective_dic = self.paper(aid)
        self.word(objective_dic, path, name)

    def quizs(self, path, name, contentId):
        payload = {
            'callCount': '1',
            'scriptSessionId': '${scriptSessionId}190',
            'httpSessionId': self.get_NTESSTUDYSI(),
            'c0-scriptName': 'MocQuizBean',
            'c0-methodName': 'getQuizInfo',
            'c0-id': '0',
            'c0-param0': f'number:{contentId}',
            'c0-param1': 'null:null',
            'c0-param2': 'boolean:false',
            'batchId': self.get_timestamp(),
        }
        url = "https://www.icourse163.org/dwr/call/plaincall/MocQuizBean.getQuizInfo.dwr"
        response = self.session.post(url=url, data=payload)
        text = response.text
        aidList = re.findall("aid=(\d+);", text)
        i = 0
        for aid in aidList:
            try:
                objective_dic = self.paper(aid)
                self.word(objective_dic, path, name, quizsbool=True, index=i)
                i = + 1
            except:
                pass

    def file_judgment(self, file, path):
        file_list = os.listdir(path)
        if file in file_list:
            return True
        else:
            return False

    def spider(self, course_id):
        Data = self.get_info_and_handle_info(course_id)
        for chaptername, chapter in Data.items():
            path = f"{self.basePath}\{chaptername}"
            if os.path.exists(path) == False:
                os.mkdir(path)
            contentType = chapter["contentType"]
            if contentType == 1:
                del chapter["contentType"]
                for lessonname, lesson in chapter.items():
                    lessonnameList = lesson.keys()
                    # 单元测试  contentType==2
                    if "contentType" in lessonnameList:
                        name = lessonname
                        contentId = lesson["contentId"]
                        b = self.file_judgment(f"{name}.docx", path)
                        if self.paperbool == True and b == False:
                            self.quizs(path, name, contentId)

                    else:
                        for unitname, unit in lesson.items():
                            contentType = unit["contentType"]
                            name = unit["name"]
                            path = f"{self.basePath}\{chaptername}\{lessonname}"
                            if os.path.exists(path) == False:
                                os.mkdir(path)
                            id = unit["id"]
                            # 视频
                            if contentType == 1:
                                b = self.file_judgment(f"{name}.mp4", path)
                                if self.moviebool and b == False:
                                    self.movie(id, path, name)

                            # 课件
                            elif contentType == 3:
                                b = self.file_judgment(f"{name}.pdf", path)
                                contentId = unit["contentId"]
                                if self.databool and b == False:
                                    self.data(path, name, id, contentId)

                            # 富文本---》目前已知富文本为思维导图(图片形式)，如有其它再添加
                            elif contentType == 4:
                                contentId = unit["contentId"]
                                if self.databool:
                                    self.text(path, name, id, contentId)

                            # 小测验
                            elif contentType == 5:
                                b = self.file_judgment(f"{name}.docx", path)
                                contentId = unit["contentId"]
                                if self.paperbool and b == False:
                                    self.test(path, name, id, contentId)

                            # 讨论
                            elif contentType == 6:
                                b = self.file_judgment(f"{name}.txt", path)
                                contentId = unit["contentId"]
                                if self.discussbool and b == False:
                                    self.discuss(path, name, id, contentId)
                            else:
                                pass
            else:
                pass

    def download_main(self, basePath, course_id, quality=3, moviebool=True, databool=True, paperbool=True,
                      discussbool=True):
        self.basePath = basePath
        self.get_csrfKey_userId()
        self.quality = quality
        self.moviebool = moviebool
        self.databool = databool
        self.paperbool = paperbool
        self.discussbool = discussbool
        self.spider(course_id)
